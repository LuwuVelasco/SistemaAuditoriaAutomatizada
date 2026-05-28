"""
Servicio de generación de reportes.
Produce XLSX o DOCX a partir de los hallazgos de una auditoría.
Sube el archivo a Supabase Storage y registra el reporte en Firestore.
"""

import io
from copy import deepcopy
from pathlib import Path
from typing import List

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from loguru import logger
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

from app.core.exceptions import ForbiddenError, NotFoundError
from app.models.audit import Audit, DEFAULT_MATURITY
from app.models.finding import Finding
from app.models.report import Report
from app.repositories.audit_repository import AuditRepository
from app.repositories.finding_repository import FindingRepository
from app.repositories.report_repository import ReportRepository
from app.schemas.report import ReportGenerateRequest
from app.services.storage_service import StorageService
from app.utils.enums import ReportFormat, ReportKind
from app.utils.helpers import generate_id
from app.utils.timestamps import utcnow_iso

_TEMPLATES_DIR = Path(__file__).resolve().parent.parent.parent / "templates"
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XML = "http://www.w3.org/XML/1998/namespace"

_CONTENT_TYPES = {
    ReportFormat.XLSX: "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ReportFormat.DOCX: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ReportFormat.PDF:  "application/pdf",
}

_EXTENSIONS = {
    ReportFormat.XLSX: "xlsx",
    ReportFormat.DOCX: "docx",
    ReportFormat.PDF:  "pdf",
}

_HEADER_FILL = PatternFill("solid", fgColor="0D1117")
_HEADER_FONT = Font(bold=True, color="22D3EE", name="Calibri", size=11)
_RISK_COLORS = {"Bajo": "22C55E", "Medio": "F59E0B", "Alto": "F97316", "Extremo": "EF4444"}

_KIND_LABELS = {
    ReportKind.MATRIZ_HALLAZGOS: "MATRIZ DE HALLAZGOS",
    ReportKind.FICHAS_HALLAZGO:  "FICHAS DE HALLAZGO",
    ReportKind.FICHAS_PRUEBAS:   "FICHAS DE PRUEBAS",
    ReportKind.MATRIZ_COSO:      "MATRIZ COSO",
    ReportKind.INFORME_MADUREZ:  "INFORME DE MADUREZ DOCUMENTAL",
}


class ReportService:
    def __init__(
        self,
        audit_repo: AuditRepository,
        finding_repo: FindingRepository,
        report_repo: ReportRepository,
        storage: StorageService,
    ):
        self._audits   = audit_repo
        self._findings = finding_repo
        self._reports  = report_repo
        self._storage  = storage

    async def generate(
        self,
        audit_id: str,
        owner_id: str,
        request: ReportGenerateRequest,
    ) -> List[Report]:
        audit = await self._audits.get_by_id(audit_id)
        if audit is None:
            raise NotFoundError("Auditoría", audit_id)
        if audit.owner_id != owner_id:
            raise ForbiddenError("No tienes acceso a esta auditoría.")

        findings = await self._findings.list_by_audit(audit_id)

        created: List[Report] = []
        for kind in request.kinds:
            report = await self._generate_one(audit, kind, request.format, findings)
            created.append(report)

        return created

    async def get_by_id(self, audit_id: str, report_id: str, owner_id: str) -> Report:
        audit = await self._audits.get_by_id(audit_id)
        if audit is None:
            raise NotFoundError("Auditoría", audit_id)
        if audit.owner_id != owner_id:
            raise ForbiddenError("No tienes acceso a esta auditoría.")
        report = await self._reports.get_by_id(audit_id, report_id)
        if report is None:
            raise NotFoundError("Reporte", report_id)
        return report

    async def list_by_audit(self, audit_id: str, owner_id: str) -> List[Report]:
        audit = await self._audits.get_by_id(audit_id)
        if audit is None:
            raise NotFoundError("Auditoría", audit_id)
        if audit.owner_id != owner_id:
            raise ForbiddenError("No tienes acceso a esta auditoría.")
        return await self._reports.list_by_audit(audit_id)

    async def _generate_one(
        self,
        audit: Audit,
        kind: ReportKind,
        fmt: ReportFormat,
        findings: List[Finding],
    ) -> Report:
        now = utcnow_iso()
        report_id = generate_id("RPT-")
        filename = f"{kind.value}_{report_id}.{_EXTENSIONS[fmt]}"

        if fmt == ReportFormat.XLSX:
            content = self._build_xlsx(kind, audit, findings)
        elif fmt == ReportFormat.DOCX:
            content = self._build_docx(kind, audit, findings)
        else:
            content = self._build_xlsx(kind, audit, findings)
            filename = filename.replace(".pdf", ".xlsx")
            fmt = ReportFormat.XLSX

        sha = StorageService.compute_sha256(content)
        path = await self._storage.upload_report(audit.id, filename, content, _CONTENT_TYPES[fmt])

        report = Report(
            id=report_id,
            auditId=audit.id,
            kind=kind,
            format=fmt,
            supabasePath=path,
            sha256=sha,
            generatedAt=now,
        )
        result = await self._reports.create(report)
        logger.info(f"Reporte generado: {kind.value} ({fmt.value}) → {path}")
        return result

    # ── XLSX builders ─────────────────────────────────────────────────────────

    def _build_xlsx(self, kind: ReportKind, audit: Audit, findings: List[Finding]) -> bytes:
        wb = Workbook()
        ws = wb.active
        ws.title = _KIND_LABELS.get(kind, kind.value)[:31]

        if kind == ReportKind.MATRIZ_HALLAZGOS:
            self._xlsx_matriz_hallazgos(ws, audit, findings)
        elif kind == ReportKind.FICHAS_HALLAZGO:
            self._xlsx_fichas_hallazgo(ws, audit, findings)
        elif kind == ReportKind.FICHAS_PRUEBAS:
            self._xlsx_fichas_pruebas(ws, audit, findings)
        elif kind == ReportKind.MATRIZ_COSO:
            self._xlsx_matriz_coso(ws, audit, findings)
        elif kind == ReportKind.INFORME_MADUREZ:
            self._xlsx_informe_madurez(ws, audit, findings)

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def _write_header_row(self, ws, headers: list, row: int = 1) -> None:
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=h)
            cell.fill = _HEADER_FILL
            cell.font = _HEADER_FONT
            cell.alignment = Alignment(horizontal="center", wrap_text=True)

    def _xlsx_matriz_hallazgos(self, ws, audit: Audit, findings: List[Finding]) -> None:
        label = _KIND_LABELS[ReportKind.MATRIZ_HALLAZGOS]
        ws["A1"] = f"{label} — {audit.entity}  |  {audit.period}  |  {audit.city}"
        ws["A1"].font = Font(bold=True, size=13, name="Calibri")
        ws.merge_cells("A1:N1")

        if audit.alcance:
            ws["A2"] = f"Alcance: {audit.alcance}"
            ws["A2"].font = Font(italic=True, size=10, name="Calibri")
            ws.merge_cells("A2:N2")
            header_row = 3
        else:
            header_row = 2

        headers = [
            "ID", "Título", "Descripción", "Criterio", "Causa", "Efecto", "Conclusión",
            "Recomendación", "Riesgo", "Impacto", "Probabilidad", "Estado", "Confianza IA", "Detectado por",
        ]
        self._write_header_row(ws, headers, row=header_row)

        for i, f in enumerate(findings, start=header_row + 1):
            risk_str = f.risk.value if hasattr(f.risk, "value") else str(f.risk)
            risk_color = _RISK_COLORS.get(risk_str, "FFFFFF")
            row_data = [
                f.id,
                f.title,
                f.description_finding or f.description,
                f.criteria_description or "",
                f.cause or "",
                f.effect or "",
                f.conclusion or "",
                f.recommendation or "",
                risk_str,
                f.impact,
                f.probability,
                f.status.value if hasattr(f.status, "value") else str(f.status),
                f"{f.confidence:.0%}",
                f.detected_by or "",
            ]
            for col, val in enumerate(row_data, 1):
                cell = ws.cell(row=i, column=col, value=val)
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                if col == 9:
                    cell.fill = PatternFill("solid", fgColor=risk_color)
                    cell.font = Font(bold=True, color="0D1117")

        ws.column_dimensions["A"].width = 14
        ws.column_dimensions["B"].width = 30
        ws.column_dimensions["C"].width = 40
        ws.column_dimensions["D"].width = 38
        ws.column_dimensions["E"].width = 28
        ws.column_dimensions["F"].width = 28
        ws.column_dimensions["G"].width = 40
        ws.column_dimensions["H"].width = 55
        ws.column_dimensions["I"].width = 12
        ws.column_dimensions["J"].width = 10
        ws.column_dimensions["K"].width = 10
        ws.column_dimensions["L"].width = 12
        ws.column_dimensions["M"].width = 14
        ws.column_dimensions["N"].width = 14

    def _xlsx_fichas_hallazgo(self, ws, audit: Audit, findings: List[Finding]) -> None:
        label = _KIND_LABELS[ReportKind.FICHAS_HALLAZGO]
        ws["A1"] = f"{label} — {audit.entity}  |  {audit.period}  |  {audit.city}"
        ws["A1"].font = Font(bold=True, size=13, name="Calibri")
        ws.merge_cells("A1:L1")

        if audit.alcance:
            ws["A2"] = f"Alcance: {audit.alcance}"
            ws["A2"].font = Font(italic=True, size=10, name="Calibri")
            ws.merge_cells("A2:L2")
            header_row = 3
        else:
            header_row = 2

        headers = [
            "ID", "Título", "Descripción", "Criterio", "Causa", "Efecto", "Conclusión",
            "Riesgo", "Recomendación", "Refs COBIT", "Refs COSO", "Refs RGSI",
        ]
        self._write_header_row(ws, headers, row=header_row)

        for i, f in enumerate(findings, start=header_row + 1):
            cobit = "; ".join(r.code for r in (f.cobit_refs or []))
            coso  = "; ".join(r.code for r in (f.coso_refs or []))
            rgsi  = "; ".join(r.code for r in (f.rgsi_refs or []))
            row_data = [
                f.id, f.title,
                f.description_finding or f.description,
                f.criteria_description or "",
                f.cause or "",
                f.effect or "",
                f.conclusion or "",
                f.risk.value if hasattr(f.risk, "value") else str(f.risk),
                f.recommendation,
                cobit, coso, rgsi,
            ]
            for col, val in enumerate(row_data, 1):
                ws.cell(row=i, column=col, value=val).alignment = Alignment(wrap_text=True, vertical="top")

        ws.column_dimensions["B"].width = 26
        ws.column_dimensions["C"].width = 36
        ws.column_dimensions["D"].width = 36
        ws.column_dimensions["E"].width = 28
        ws.column_dimensions["F"].width = 28
        ws.column_dimensions["G"].width = 36
        ws.column_dimensions["I"].width = 60

    def _xlsx_fichas_pruebas(self, ws, audit: Audit, findings: List[Finding]) -> None:
        label = _KIND_LABELS[ReportKind.FICHAS_PRUEBAS]
        ws["A1"] = f"{label} — {audit.entity}  |  {audit.period}"
        ws["A1"].font = Font(bold=True, size=13, name="Calibri")
        ws.merge_cells("A1:H1")

        if audit.alcance:
            ws["A2"] = f"Alcance: {audit.alcance}"
            ws["A2"].font = Font(italic=True, size=10, name="Calibri")
            ws.merge_cells("A2:H2")
            header_row = 3
        else:
            header_row = 2

        headers = [
            "ID Hallazgo", "Título Hallazgo",
            "Nombre Documento", "ID Documento", "Página", "Párrafo",
            "Cita", "Estado",
        ]
        self._write_header_row(ws, headers, row=header_row)

        row_idx = header_row + 1
        for f in findings:
            evidence_list = f.evidence or []
            status_val = f.status.value if hasattr(f.status, "value") else str(f.status)
            quote_val = f.quote or ""

            if not evidence_list:
                row_data = [f.id, f.title, "—", "—", "", "", quote_val, status_val]
                for col, val in enumerate(row_data, 1):
                    ws.cell(row=row_idx, column=col, value=val).alignment = Alignment(wrap_text=True, vertical="top")
                row_idx += 1
            else:
                for ev in evidence_list:
                    row_data = [
                        f.id, f.title,
                        ev.doc_name,
                        ev.doc_id,
                        ev.page if ev.page is not None else "",
                        ev.paragraph or "",
                        quote_val,
                        status_val,
                    ]
                    for col, val in enumerate(row_data, 1):
                        ws.cell(row=row_idx, column=col, value=val).alignment = Alignment(wrap_text=True, vertical="top")
                    row_idx += 1

        ws.column_dimensions["A"].width = 14
        ws.column_dimensions["B"].width = 35
        ws.column_dimensions["C"].width = 35
        ws.column_dimensions["D"].width = 22
        ws.column_dimensions["E"].width = 8
        ws.column_dimensions["F"].width = 50
        ws.column_dimensions["G"].width = 50
        ws.column_dimensions["H"].width = 14

    def _xlsx_matriz_coso(self, ws, audit: Audit, findings: List[Finding]) -> None:
        label = _KIND_LABELS[ReportKind.MATRIZ_COSO]
        ws["A1"] = f"{label} — {audit.entity}  |  {audit.period}"
        ws["A1"].font = Font(bold=True, size=13, name="Calibri")
        ws.merge_cells("A1:G1")

        if audit.alcance:
            ws["A2"] = f"Alcance: {audit.alcance}"
            ws["A2"].font = Font(italic=True, size=10, name="Calibri")
            ws.merge_cells("A2:G2")
            header_row = 3
        else:
            header_row = 2

        headers = ["ID", "Título", "Componente COSO", "Principio", "Riesgo", "Recomendación", "Estado"]
        self._write_header_row(ws, headers, row=header_row)

        coso_findings = [f for f in findings if f.coso_refs]

        for i, f in enumerate(coso_findings, start=header_row + 1):
            coso_ref = f.coso_refs or []
            comp  = coso_ref[0].component if coso_ref and hasattr(coso_ref[0], "component") else ""
            princ = coso_ref[0].title if coso_ref else ""
            row_data = [
                f.id, f.title, comp or "", princ,
                f.risk.value if hasattr(f.risk, "value") else str(f.risk),
                f.recommendation,
                f.status.value if hasattr(f.status, "value") else str(f.status),
            ]
            for col, val in enumerate(row_data, 1):
                ws.cell(row=i, column=col, value=val).alignment = Alignment(wrap_text=True, vertical="top")

        ws.column_dimensions["B"].width = 35
        ws.column_dimensions["F"].width = 55

    # ── DOCX builder (dispatch) ───────────────────────────────────────────────

    def _build_docx(self, kind: ReportKind, audit: Audit, findings: List[Finding]) -> bytes:
        if kind == ReportKind.FICHAS_PRUEBAS:
            return self._docx_fichas_pruebas(audit, findings)
        if kind == ReportKind.FICHAS_HALLAZGO:
            return self._docx_fichas_hallazgo(audit, findings)
        if kind == ReportKind.INFORME_MADUREZ:
            return self._docx_informe_madurez(audit, findings)
        return self._docx_generic(kind, audit, findings)

    # ── DOCX: Fichas de Pruebas (basado en template 4. Pruebas.docx) ─────────

    def _docx_fichas_pruebas(self, audit: Audit, findings: List[Finding]) -> bytes:
        template_path = _TEMPLATES_DIR / "4. Pruebas.docx"

        try:
            doc = DocxDocument(str(template_path))
        except Exception as exc:
            logger.warning(f"No se pudo cargar template pruebas: {exc}. Usando generación genérica.")
            return self._docx_generic(ReportKind.FICHAS_PRUEBAS, audit, findings)

        if not doc.tables:
            return self._docx_generic(ReportKind.FICHAS_PRUEBAS, audit, findings)

        # Guardar XML de referencia del primer bloque de prueba y su fila de evidencia
        ref_tbl_xml = deepcopy(doc.tables[0]._element)
        ref_rows = ref_tbl_xml.findall(f"{{{_W}}}tr")
        ref_ev_row = deepcopy(ref_rows[7]) if len(ref_rows) > 7 else None

        # Limpiar cuerpo del documento (preserva estilos y metadatos del template)
        body = doc.element.body
        for el in list(body):
            body.remove(el)

        # Encabezado del documento
        self._docx_para(doc, f"BANCO AUDITADO: {audit.entity}", bold=True)
        self._docx_para(doc, f"Período: {audit.period}  |  Ciudad: {audit.city}")
        if audit.alcance:
            self._docx_para(doc, f"Alcance: {audit.alcance}", italic=True)
        self._docx_para(doc, "IDENTIFICACIÓN DE EVIDENCIAS EN BASE A LOS HALLAZGOS", bold=True)
        self._docx_para(doc, "GUÍA DE PRUEBAS", bold=True)

        for f in findings:
            self._docx_para(doc, "")  # espacio entre tablas

            tbl_xml = deepcopy(ref_tbl_xml)
            rows = tbl_xml.findall(f"{{{_W}}}tr")

            # Estructura real del template (celdas con gridSpan):
            # Row 1: tc[0]=span3 "Guía de Pruebas", tc[1]=span1 → ID hallazgo
            # Row 2-5: tc[0]=span2 label, tc[1]=span2 → valor
            # Row 7+: tc[0]="No", tc[1]="Evidencia", tc[2]=span2 → Descripción

            self._xml_set(rows, 1, 1, f.id)

            domain = ""
            if f.cobit_refs:
                domain = getattr(f.cobit_refs[0], "domain", None) or f.cobit_refs[0].title
            elif f.coso_refs:
                domain = getattr(f.coso_refs[0], "component", None) or f.coso_refs[0].title
            elif f.rgsi_refs:
                domain = getattr(f.rgsi_refs[0], "section", None) or f.rgsi_refs[0].title
            self._xml_set(rows, 2, 1, domain)

            all_refs = f.cobit_refs or f.coso_refs or f.rgsi_refs
            process = "\n".join(f"{r.code} {r.title}" for r in all_refs) if all_refs else ""
            self._xml_set(rows, 3, 1, process)

            objective = f.criteria_description or f.description_finding or f.description
            self._xml_set(rows, 4, 1, objective)

            risk_str = f.risk.value if hasattr(f.risk, "value") else str(f.risk)
            self._xml_set(rows, 5, 1, risk_str)

            # Eliminar filas de evidencia existentes (fila 7 en adelante)
            for ev_row_el in rows[7:]:
                tbl_xml.remove(ev_row_el)

            # Agregar filas de evidencia reales
            evidence_list = f.evidence or []
            if not evidence_list or ref_ev_row is None:
                if ref_ev_row is not None:
                    new_row = deepcopy(ref_ev_row)
                    # tc[0]=No, tc[1]=Evidencia, tc[2]=Descripción (span2)
                    self._xml_set_row(new_row, 0, "1")
                    self._xml_set_row(new_row, 1, "—")
                    desc = f.description_finding or f.description or ""
                    self._xml_set_row(new_row, 2, desc)
                    tbl_xml.append(new_row)
            else:
                for ev_n, ev in enumerate(evidence_list, 1):
                    new_row = deepcopy(ref_ev_row)
                    ev_label = ev.doc_name
                    if ev.page is not None:
                        ev_label += f", Pág. {ev.page}"
                    if ev.paragraph and len(ev.paragraph) < 80:
                        ev_label += f" — {ev.paragraph}"
                    desc = ev.paragraph or f.description_finding or f.description or ""
                    # tc[0]=No, tc[1]=Evidencia, tc[2]=Descripción (span2)
                    self._xml_set_row(new_row, 0, str(ev_n))
                    self._xml_set_row(new_row, 1, ev_label)
                    self._xml_set_row(new_row, 2, desc)
                    tbl_xml.append(new_row)

            body.append(tbl_xml)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── DOCX: Fichas de Hallazgo (basado en template 3. fichas de hallazgos.docx) ─

    def _docx_fichas_hallazgo(self, audit: Audit, findings: List[Finding]) -> bytes:
        template_path = _TEMPLATES_DIR / "3. fichas de hallazgos.docx"

        try:
            doc = DocxDocument(str(template_path))
        except Exception as exc:
            logger.warning(f"No se pudo cargar template fichas hallazgo: {exc}. Usando generación genérica.")
            return self._docx_generic(ReportKind.FICHAS_HALLAZGO, audit, findings)

        if not doc.tables:
            return self._docx_generic(ReportKind.FICHAS_HALLAZGO, audit, findings)

        # Reference table XML — cloned for each finding × framework
        ref_tbl_xml = deepcopy(doc.tables[0]._element)

        body = doc.element.body
        for el in list(body):
            body.remove(el)

        self._docx_para(doc, f"BANCO AUDITADO: {audit.entity}", bold=True)
        self._docx_para(doc, f"Período: {audit.period}  |  Ciudad: {audit.city}")
        if audit.alcance:
            self._docx_para(doc, f"Alcance: {audit.alcance}", italic=True)
        self._docx_para(doc, "FICHAS DE HALLAZGO", bold=True)

        # (fw_name, row2_label, row3_label, refs_attr)
        fw_configs = [
            ("COBIT", "Dominio",    "Proceso",      "cobit_refs"),
            ("COSO",  "Componente", "Lineamiento",  "coso_refs"),
            ("RGSI",  "Sección",    "Artículos",    "rgsi_refs"),
        ]

        for f in findings:
            risk_str = f.risk.value if hasattr(f.risk, "value") else str(f.risk)

            for fw_name, label2, label3, refs_attr in fw_configs:
                refs = getattr(f, refs_attr, None) or []
                if not refs:
                    continue

                self._docx_para(doc, "")

                tbl_xml = deepcopy(ref_tbl_xml)
                rows = tbl_xml.findall(f"{{{_W}}}tr")

                # Remove "Riesgos Asociados" row (index 5) — not needed per requirement
                if len(rows) > 5:
                    tbl_xml.remove(rows[5])
                    rows = tbl_xml.findall(f"{{{_W}}}tr")

                # Row 1 tc[1] → Finding ID
                self._xml_set(rows, 1, 1, f.id)

                # Row 2: tc[0] = framework label, tc[1] = domain/component/section value
                self._xml_set(rows, 2, 0, label2)
                if fw_name == "COBIT":
                    val2 = getattr(refs[0], "domain", None) or refs[0].title
                elif fw_name == "COSO":
                    val2 = getattr(refs[0], "component", None) or refs[0].title
                else:
                    val2 = getattr(refs[0], "section", None) or refs[0].title
                self._xml_set(rows, 2, 1, val2 or "")

                # Row 3: tc[0] = sub-label, tc[1] = refs text
                self._xml_set(rows, 3, 0, label3)
                refs_text = "\n".join(f"{r.code} — {r.title}" for r in refs)
                self._xml_set(rows, 3, 1, refs_text)

                # Row 4: tc[1] → criteria / objective
                self._xml_set(rows, 4, 1, f.criteria_description or f.description_finding or "")

                # After row 5 removed: old rows 6-13 become 5-12
                # Rows 5/7/9/11 are labels (span3); rows 6/8/10/12 are content (span3)
                self._xml_set(rows, 6,  0, f.description_finding or f.description or "")
                self._xml_set(rows, 8,  0, f.recommendation or "")
                self._xml_set(rows, 10, 0, f.cause or "")
                self._xml_set(rows, 12, 0, risk_str)

                body.append(tbl_xml)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── DOCX: Genérico (para otros tipos de reporte) ──────────────────────────

    def _docx_generic(self, kind: ReportKind, audit: Audit, findings: List[Finding]) -> bytes:
        doc = DocxDocument()
        label = _KIND_LABELS.get(kind, kind.value.upper().replace("-", " "))

        title = doc.add_heading(f"{label} — {audit.entity}", 0)
        title.runs[0].font.color.rgb = RGBColor(0x22, 0xD3, 0xEE)

        meta_tbl = doc.add_table(rows=1, cols=2)
        meta_tbl.style = "Table Grid"
        self._docx_row(meta_tbl, "Entidad", audit.entity)
        self._docx_row(meta_tbl, "Tipo", audit.type)
        self._docx_row(meta_tbl, "Período", audit.period)
        self._docx_row(meta_tbl, "Ciudad", audit.city)
        if audit.alcance:
            self._docx_row(meta_tbl, "Alcance", audit.alcance)
        self._docx_row(meta_tbl, "Total hallazgos", str(len(findings)))
        doc.add_paragraph("")

        report_findings = [f for f in findings if f.coso_refs] if kind == ReportKind.MATRIZ_COSO else findings

        for f in report_findings:
            h = doc.add_heading(f"[{f.id}] {f.title}", level=2)
            risk_str = f.risk.value if hasattr(f.risk, "value") else str(f.risk)
            h.runs[0].font.color.rgb = RGBColor(0xF9, 0x7B, 0x16) if risk_str in ("Alto", "Extremo") else RGBColor(0x22, 0xD3, 0xEE)

            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = "Table Grid"
            self._docx_row(tbl, "Descripción", f.description_finding or f.description)
            self._docx_row(tbl, "Criterio", f.criteria_description or "")
            self._docx_row(tbl, "Causa", f.cause or "")
            self._docx_row(tbl, "Efecto", f.effect or "")
            self._docx_row(tbl, "Conclusión", f.conclusion or "")
            self._docx_row(tbl, "Riesgo", risk_str)
            self._docx_row(tbl, "Estado", f.status.value if hasattr(f.status, "value") else str(f.status))
            self._docx_row(tbl, "Impacto", str(f.impact))
            self._docx_row(tbl, "Probabilidad", str(f.probability))
            self._docx_row(tbl, "Confianza IA", f"{f.confidence:.0%}")
            doc.add_paragraph("")

            doc.add_paragraph("Recomendación")
            p2 = doc.add_paragraph(f.recommendation)
            p2.runs[0].font.size = Pt(10)

            if f.quote:
                doc.add_paragraph("Cita del documento")
                q = doc.add_paragraph(f'"{f.quote}"')
                q.runs[0].italic = True
                q.runs[0].font.size = Pt(9)

            doc.add_paragraph("─" * 80)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── Helpers XML (para modificar celdas del template clonado) ─────────────

    @staticmethod
    def _xml_set(rows: list, row_idx: int, col_idx: int, text: str) -> None:
        """Setea texto en una celda identificada por índice de fila y columna."""
        if row_idx >= len(rows):
            return
        cells = rows[row_idx].findall(f"{{{_W}}}tc")
        if col_idx >= len(cells):
            return
        ReportService._xml_write_cell(cells[col_idx], text)

    @staticmethod
    def _xml_set_row(row_el, col_idx: int, text: str) -> None:
        """Setea texto en una celda de una fila XML dada."""
        cells = row_el.findall(f"{{{_W}}}tc")
        if col_idx >= len(cells):
            return
        ReportService._xml_write_cell(cells[col_idx], text)

    @staticmethod
    def _xml_write_cell(cell_el, text: str) -> None:
        """Reemplaza el contenido de texto en una celda preservando estilos de párrafo."""
        paragraphs = cell_el.findall(f"{{{_W}}}p")
        if not paragraphs:
            return

        lines = (text or "").split("\n")
        p0 = paragraphs[0]

        # Limpiar runs del primer párrafo preservando propiedades (pPr, rPr)
        for run in list(p0.findall(f"{{{_W}}}r")):
            p0.remove(run)

        # Escribir primera línea en el párrafo existente
        if lines[0]:
            from lxml.etree import SubElement
            r_el = SubElement(p0, f"{{{_W}}}r")
            t_el = SubElement(r_el, f"{{{_W}}}t")
            t_el.text = lines[0]
            if lines[0].startswith(" ") or lines[0].endswith(" "):
                t_el.set(f"{{{_XML}}}space", "preserve")

        # Eliminar párrafos adicionales originales
        for extra_p in paragraphs[1:]:
            cell_el.remove(extra_p)

        # Agregar párrafos para líneas adicionales
        for line in lines[1:]:
            new_p = deepcopy(p0)
            for run in list(new_p.findall(f"{{{_W}}}r")):
                new_p.remove(run)
            if line:
                from lxml.etree import SubElement
                r_el = SubElement(new_p, f"{{{_W}}}r")
                t_el = SubElement(r_el, f"{{{_W}}}t")
                t_el.text = line
            cell_el.append(new_p)

    # ── Helpers DOCX de alto nivel ────────────────────────────────────────────

    @staticmethod
    def _docx_para(doc: DocxDocument, text: str, bold: bool = False, italic: bool = False) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic

    @staticmethod
    def _docx_row(table, label: str, value: str) -> None:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value

    def _xlsx_informe_madurez(self, ws, audit: Audit, findings: List[Finding]) -> None:
        label = _KIND_LABELS[ReportKind.INFORME_MADUREZ]
        ws["A1"] = f"{label} — {audit.entity}  |  {audit.period}"
        ws["A1"].font = Font(bold=True, size=13, name="Calibri")
        ws.merge_cells("A1:G1")

        maturity = getattr(audit, "maturity", None) or DEFAULT_MATURITY
        level = maturity.get("level", 1)
        scores = maturity.get("scores", {})
        checklist = maturity.get("checklist", {})
        gap = maturity.get("gapAnalysis", {})

        levels_map = {
            1: "Nivel 1: Acumulación Digital (Inicial)",
            2: "Nivel 2: Normas Básicas, Aplicación Desigual (Controlado)",
            3: "Nivel 3: Procesos Alineados (Integrado)",
            4: "Nivel 4: Gobierno de la Información (Estratégico)",
            5: "Nivel 5: Cultura Documental (Optimizado)"
        }
        level_text = levels_map.get(level, f"Nivel {level}")

        # Escribir información general
        ws["A3"] = "EVALUACIÓN EJECUTIVA"
        ws["A3"].font = Font(bold=True, size=11, name="Calibri", color="22D3EE")
        ws["B3"] = level_text
        ws["B3"].font = Font(bold=True, name="Calibri")

        # Escribir puntajes por dimensión
        ws["A5"] = "DIMENSIONES"
        ws["A5"].font = Font(bold=True, name="Calibri")
        ws["B5"] = "PUNTAJE (%)"
        ws["B5"].font = Font(bold=True, name="Calibri")

        dims = [
            ("Políticas y Normativas", scores.get("policies", 0)),
            ("Procesos y Roles", scores.get("processes", 0)),
            ("Trazabilidad y Control", scores.get("traceability", 0)),
            ("Cultura y Cumplimiento", scores.get("culture", 0))
        ]
        for row_idx, (dim_name, val) in enumerate(dims, start=6):
            ws.cell(row=row_idx, column=1, value=dim_name).font = Font(name="Calibri")
            ws.cell(row=row_idx, column=2, value=f"{val}%").font = Font(name="Calibri")

        # Escribir checklist
        ws["A11"] = "CHECKLIST DE EVALUACIÓN"
        ws["A11"].font = Font(bold=True, size=11, name="Calibri", color="22D3EE")
        ws.merge_cells("A11:C11")

        headers = ["Indicador", "Nivel", "Estado"]
        for col_idx, h in enumerate(headers, start=1):
            cell = ws.cell(row=12, column=col_idx, value=h)
            cell.font = Font(bold=True, name="Calibri")

        checklist_items = [
            ("l1_repositorios", "Existen repositorios electrónicos para almacenar los documentos.", "L1", "Requerido"),
            ("l1_sin_politicas", "No hay políticas documentales formalmente aprobadas.", "L1", "Riesgo"),
            ("l2_cuadro", "Existe un cuadro de clasificación documental aprobado.", "L2", "Requerido"),
            ("l2_calendario", "Existen calendarios de conservación de documentos.", "L2", "Requerido"),
            ("l2_desigual", "La aplicación de criterios archivísticos es desigual en las áreas.", "L2", "Riesgo"),
            ("l3_procesos", "La gestión documental está integrada a los procedimientos operativos.", "L3", "Requerido"),
            ("l3_roles", "Existen roles claros definidos (técnicos, funcionales y jurídicos).", "L3", "Requerido"),
            ("l3_trazabilidad", "Hay trazabilidad documental real y ciclo de vida conocido.", "L3", "Requerido"),
            ("l4_riesgos", "La dirección conoce activamente los riesgos documentales.", "L4", "Requerido"),
            ("l4_activos", "La información y los documentos se gestionan como activos corporativos.", "L4", "Requerido"),
            ("l4_coordinacion", "Existe coordinación efectiva entre IT, jurídico y compliance.", "L4", "Requerido"),
            ("l5_cultura", "La gestión documental es parte intrínseca de la cultura corporativa.", "L5", "Requerido"),
            ("l5_mejora", "Existen procesos activos de mejora continua en la gestión documental.", "L5", "Requerido"),
            ("l5_inspeccion", "La información está permanentemente preparada para auditorías o litigios.", "L5", "Requerido")
        ]

        curr_row = 13
        for key, desc, lvl, kind_item in checklist_items:
            status = "CUMPLE" if checklist.get(key, False) else "NO CUMPLE"
            if kind_item == "Riesgo":
                status = "PRESENTE (Riesgo)" if checklist.get(key, False) else "AUSENTE (Controlado)"
            
            ws.cell(row=curr_row, column=1, value=desc).font = Font(name="Calibri")
            ws.cell(row=curr_row, column=2, value=lvl).font = Font(name="Calibri")
            
            cell_status = ws.cell(row=curr_row, column=3, value=status)
            cell_status.font = Font(bold=True, name="Calibri")
            if "NO CUMPLE" in status or "PRESENTE" in status:
                cell_status.font = Font(bold=True, color="EF4444", name="Calibri")
            else:
                cell_status.font = Font(bold=True, color="22C55E", name="Calibri")
            curr_row += 1

        # Análisis de brechas
        curr_row += 2
        ws.cell(row=curr_row, column=1, value="ANÁLISIS DE BRECHAS (GAP ANALYSIS)").font = Font(bold=True, size=11, name="Calibri", color="22D3EE")
        ws.merge_cells(f"A{curr_row}:C{curr_row}")
        
        sections = [
            ("Puntos Fuertes", gap.get("strengths", [])),
            ("Brechas Clave / Debilidades", gap.get("weaknesses", [])),
            ("Plan de Acción / Roadmap", gap.get("roadmap", []))
        ]
        
        for title, items in sections:
            curr_row += 1
            ws.cell(row=curr_row, column=1, value=title.upper()).font = Font(bold=True, name="Calibri")
            for item in items:
                curr_row += 1
                ws.cell(row=curr_row, column=1, value=f"• {item}").font = Font(name="Calibri")

        ws.column_dimensions["A"].width = 75
        ws.column_dimensions["B"].width = 25
        ws.column_dimensions["C"].width = 25

    def _docx_informe_madurez(self, audit: Audit, findings: List[Finding]) -> bytes:
        doc = DocxDocument()
        label = _KIND_LABELS[ReportKind.INFORME_MADUREZ]

        title = doc.add_heading(f"{label}", 0)
        title.runs[0].font.color.rgb = RGBColor(0x22, 0xD3, 0xEE)

        meta_tbl = doc.add_table(rows=1, cols=2)
        meta_tbl.style = "Table Grid"
        self._docx_row(meta_tbl, "Entidad", audit.entity)
        self._docx_row(meta_tbl, "Tipo", audit.type)
        self._docx_row(meta_tbl, "Período", audit.period)
        self._docx_row(meta_tbl, "Ciudad", audit.city)

        maturity = getattr(audit, "maturity", None) or DEFAULT_MATURITY
        level = maturity.get("level", 1)
        scores = maturity.get("scores", {})
        checklist = maturity.get("checklist", {})
        gap = maturity.get("gapAnalysis", {})

        levels_map = {
            1: "Nivel 1: Acumulación Digital (Inicial)",
            2: "Nivel 2: Normas Básicas, Aplicación Desigual (Controlado)",
            3: "Nivel 3: Procesos Alineados (Integrado)",
            4: "Nivel 4: Gobierno de la Información (Estratégico)",
            5: "Nivel 5: Cultura Documental (Optimizado)"
        }
        level_text = levels_map.get(level, f"Nivel {level}")

        doc.add_paragraph("")
        self._docx_para(doc, f"NIVEL DE MADUREZ DOCUMENTAL OBTENIDO: {level_text}", bold=True)
        doc.add_paragraph("")

        # Agregar tabla de dimensiones
        doc.add_heading("Puntuaciones por Dimensión", level=2)
        dim_tbl = doc.add_table(rows=1, cols=2)
        dim_tbl.style = "Table Grid"
        self._docx_row(dim_tbl, "Dimensión", "Puntaje Obtenido (%)")
        self._docx_row(dim_tbl, "Políticas y Normativas", f"{scores.get('policies', 0)}%")
        self._docx_row(dim_tbl, "Procesos y Roles", f"{scores.get('processes', 0)}%")
        self._docx_row(dim_tbl, "Trazabilidad y Control", f"{scores.get('traceability', 0)}%")
        self._docx_row(dim_tbl, "Cultura y Cumplimiento", f"{scores.get('culture', 0)}%")

        # Agregar checklist
        doc.add_paragraph("")
        doc.add_heading("Checklist Detallado de Autoevaluación", level=2)
        check_tbl = doc.add_table(rows=1, cols=3)
        check_tbl.style = "Table Grid"
        
        # Escribir cabecera del checklist
        hdr_cells = check_tbl.rows[0].cells
        hdr_cells[0].text = "Criterio de Evaluación"
        hdr_cells[0].paragraphs[0].runs[0].bold = True
        hdr_cells[1].text = "Nivel"
        hdr_cells[1].paragraphs[0].runs[0].bold = True
        hdr_cells[2].text = "Estado"
        hdr_cells[2].paragraphs[0].runs[0].bold = True

        checklist_items = [
            ("l1_repositorios", "Existen repositorios electrónicos para almacenar los documentos.", "L1", "Requerido"),
            ("l1_sin_politicas", "No hay políticas documentales formalmente aprobadas.", "L1", "Riesgo"),
            ("l2_cuadro", "Existe un cuadro de clasificación documental aprobado.", "L2", "Requerido"),
            ("l2_calendario", "Existen calendarios de conservación de documentos.", "L2", "Requerido"),
            ("l2_desigual", "La aplicación de criterios archivísticos es desigual en las áreas.", "L2", "Riesgo"),
            ("l3_procesos", "La gestión documental está integrada a los procedimientos operativos.", "L3", "Requerido"),
            ("l3_roles", "Existen roles claros definidos (técnicos, funcionales y jurídicos).", "L3", "Requerido"),
            ("l3_trazabilidad", "Hay trazabilidad documental real y ciclo de vida conocido.", "L3", "Requerido"),
            ("l4_riesgos", "La dirección conoce activamente los riesgos documentales.", "L4", "Requerido"),
            ("l4_activos", "La información y los documentos se gestionan como activos corporativos.", "L4", "Requerido"),
            ("l4_coordinacion", "Existe coordinación efectiva entre IT, jurídico y compliance.", "L4", "Requerido"),
            ("l5_cultura", "La gestión documental es parte intrínseca de la cultura corporativa.", "L5", "Requerido"),
            ("l5_mejora", "Existen procesos activos de mejora continua en la gestión documental.", "L5", "Requerido"),
            ("l5_inspeccion", "La información está permanentemente preparada para auditorías o litigios.", "L5", "Requerido")
        ]

        for key, desc, lvl, kind_item in checklist_items:
            status = "CUMPLE" if checklist.get(key, False) else "NO CUMPLE"
            if kind_item == "Riesgo":
                status = "PRESENTE (Riesgo)" if checklist.get(key, False) else "AUSENTE (Controlado)"
                
            row = check_tbl.add_row()
            row.cells[0].text = desc
            row.cells[1].text = lvl
            row.cells[2].text = status
            
            # Formatear estado con color
            cell_run = row.cells[2].paragraphs[0].runs[0]
            cell_run.bold = True
            if "NO CUMPLE" in status or "PRESENTE" in status:
                cell_run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
            else:
                cell_run.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)

        # Agregar brechas y plan de acción
        doc.add_paragraph("")
        doc.add_heading("Análisis de Brechas (Gap Analysis)", level=2)
        
        doc.add_heading("PUNTOS FUERTES", level=3)
        for s in gap.get("strengths", []):
            doc.add_paragraph(f"• {s}")
            
        doc.add_heading("BRECHAS CLAVE / DEBILIDADES", level=3)
        for w in gap.get("weaknesses", []):
            doc.add_paragraph(f"• {w}")
            
        doc.add_heading("PLAN DE ACCIÓN / ROADMAP SUGERIDO", level=3)
        for r in gap.get("roadmap", []):
            doc.add_paragraph(f"• {r}")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
