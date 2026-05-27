"""
Servicio de extracción de texto de documentos.

Arquitectura preparada para:
- pypdf (texto embebido)
- OCR futuro (pytesseract / Google Document AI)
- Chunking con solapamiento
- Generación de embeddings
"""

import io
from pathlib import Path
from typing import List

from loguru import logger

try:
    from pypdf import PdfReader
    _PYPDF_AVAILABLE = True
except ImportError:
    _PYPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    from openpyxl import load_workbook
    _XLSX_AVAILABLE = True
except ImportError:
    _XLSX_AVAILABLE = False

from app.utils.helpers import chunk_text, normalize_text


class ExtractionService:
    """Extrae y normaliza texto de documentos para el pipeline IA."""

    # Tamaño de chunk en caracteres (≈ 750 tokens Gemini)
    CHUNK_SIZE = 3000
    CHUNK_OVERLAP = 200

    def extract_text_from_pdf(self, content: bytes) -> str:
        """
        Extrae texto de un PDF con texto embebido.
        En el futuro: fallback a OCR si el texto extraído es vacío o muy corto.
        """
        if not _PYPDF_AVAILABLE:
            logger.warning("pypdf no disponible — texto vacío.")
            return ""

        try:
            reader = PdfReader(io.BytesIO(content))
            pages_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages_text.append(f"[Página {i + 1}]\n{text}")
            full_text = "\n\n".join(pages_text)
            logger.debug(f"PDF extraído: {len(reader.pages)} páginas, {len(full_text)} caracteres.")
            return full_text
        except Exception as exc:
            logger.error(f"Error al extraer texto del PDF: {exc}")
            return ""

    def extract_text_from_docx(self, content: bytes) -> str:
        if not _DOCX_AVAILABLE:
            logger.warning("python-docx no disponible — texto vacío.")
            return ""

        try:
            document = DocxDocument(io.BytesIO(content))
            parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        parts.append(row_text)

            full_text = "\n".join(parts)
            logger.debug(f"DOCX extraído: {len(full_text)} caracteres.")
            return full_text
        except Exception as exc:
            logger.error(f"Error al extraer texto del DOCX: {exc}")
            return ""

    def extract_text_from_xlsx(self, content: bytes) -> str:
        if not _XLSX_AVAILABLE:
            logger.warning("openpyxl no disponible — texto vacío.")
            return ""

        try:
            workbook = load_workbook(io.BytesIO(content), data_only=True)
            sheets = []
            for sheet in workbook.worksheets:
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                    if values:
                        rows.append(" | ".join(values))
                if rows:
                    sheets.append(f"[Hoja: {sheet.title}]\n" + "\n".join(rows))

            full_text = "\n\n".join(sheets)
            logger.debug(f"XLSX extraído: {len(workbook.worksheets)} hojas, {len(full_text)} caracteres.")
            return full_text
        except Exception as exc:
            logger.error(f"Error al extraer texto del XLSX: {exc}")
            return ""

    def normalize(self, text: str) -> str:
        """Normaliza el texto extraído para procesamiento IA."""
        return normalize_text(text)

    def chunk(self, text: str) -> List[str]:
        """Divide el texto en fragmentos con solapamiento."""
        return chunk_text(text, self.CHUNK_SIZE, self.CHUNK_OVERLAP)

    def process(self, content: bytes, filename: str = "") -> tuple[str, List[str]]:
        """
        Pipeline completo: extracción → normalización → chunking.
        Retorna (texto_completo, lista_de_chunks).
        """
        ext = Path(filename).suffix.lower().lstrip(".") if filename else "pdf"
        if ext == "pdf":
            raw = self.extract_text_from_pdf(content)
        elif ext == "docx":
            raw = self.extract_text_from_docx(content)
        elif ext == "xlsx":
            raw = self.extract_text_from_xlsx(content)
        else:
            logger.warning(f"Formato no soportado ({ext}) — intentando como texto PDF.")
            raw = self.extract_text_from_pdf(content)
        normalized = self.normalize(raw)
        chunks = self.chunk(normalized) if normalized else []
        logger.info(f"Extracción completa: {len(normalized)} chars → {len(chunks)} chunks.")
        return normalized, chunks
