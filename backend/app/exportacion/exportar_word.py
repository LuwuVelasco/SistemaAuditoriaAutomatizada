"""
Exportación a documentos Word usando docxtpl.
Genera las Fichas de Hallazgos y la Guía de Pruebas con variables dinámicas.
"""

from __future__ import annotations
from pathlib import Path

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..config import SALIDAS_DIR, PLANTILLAS_ORIGINALES_DIR
from ..models import Hallazgo, ContextoEntidad


def generar_fichas_hallazgos(
    hallazgos: list[Hallazgo],
    contexto: ContextoEntidad,
) -> Path | None:
    """
    Genera el documento Word de Fichas de Hallazgos.

    Crea un documento desde cero que replica la estructura de la plantilla
    original, inyectando los datos de los hallazgos encontrados.

    Args:
        hallazgos: Lista de hallazgos validados.
        contexto: Contexto de la entidad auditada.

    Returns:
        Ruta al archivo generado, o None si falla.
    """
    try:
        doc = Document()

        # Configurar estilos base
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

        # ── Encabezado ────────────────────────────────────────────
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("AUDITORES: ")
        run.bold = True
        run.font.size = Pt(11)

        grupo = contexto.grupo_auditor or "Grupo Auditor"
        p = doc.add_paragraph(grupo)
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

        if contexto.integrantes:
            p = doc.add_paragraph()
            run = p.add_run("Integrantes: ")
            run.bold = True
            for nombre in contexto.integrantes:
                doc.add_paragraph(nombre, style="List Bullet")

        p = doc.add_paragraph()
        run = p.add_run("BANCO AUDITADO: ")
        run.bold = True
        run.font.size = Pt(11)
        p = doc.add_paragraph(contexto.nombre_entidad)
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)
        p.runs[0].font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph()  # Espacio

        # Título principal
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("GUÍA DE HALLAZGOS")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph()  # Espacio

        # ── Generar una ficha por hallazgo ────────────────────────
        prefijo = _generar_prefijo_entidad(contexto.nombre_entidad)

        for i, hallazgo in enumerate(hallazgos):
            if i > 0:
                doc.add_page_break()

            codigo_ficha = f"{prefijo}-{hallazgo.codigo}"

            # Tabla de la ficha
            tabla = doc.add_table(rows=8, cols=2)
            tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
            tabla.style = "Table Grid"

            # Fila 0: Código
            _set_celda(tabla, 0, 0, "Cod.", bold=True, bg="003366", fg="FFFFFF")
            _set_celda(tabla, 0, 1, f"Hallazgos de la Auditoría\n{codigo_ficha}", bold=True)

            # Fila 1: Dominio
            _set_celda(tabla, 1, 0, "Dominio", bold=True, bg="D9E2F3")
            _set_celda(tabla, 1, 1, hallazgo.dominio)

            # Fila 2: Proceso
            _set_celda(tabla, 2, 0, "Proceso", bold=True, bg="D9E2F3")
            _set_celda(tabla, 2, 1, "\n".join(hallazgo.procesos))

            # Fila 3: Objetivo de Control
            _set_celda(tabla, 3, 0, "Objetivo de Control", bold=True, bg="D9E2F3")
            _set_celda(tabla, 3, 1, hallazgo.objetivo_control)

            # Fila 4: Riesgos Asociados
            _set_celda(tabla, 4, 0, "Riesgos Asociados", bold=True, bg="D9E2F3")
            _set_celda(tabla, 4, 1, ", ".join(hallazgo.riesgos_asociados))

            # Fila 5: Descripción
            _set_celda(tabla, 5, 0, "Descripción", bold=True, bg="D9E2F3")
            _set_celda(tabla, 5, 1, hallazgo.descripcion)

            # Fila 6: Recomendación
            _set_celda(tabla, 6, 0, "Recomendación", bold=True, bg="D9E2F3")
            _set_celda(tabla, 6, 1, hallazgo.recomendacion)

            # Fila 7: Causa
            _set_celda(tabla, 7, 0, "Causa", bold=True, bg="D9E2F3")
            _set_celda(tabla, 7, 1, hallazgo.causa)

            # Agregar fila de Nivel del Riesgo
            row = tabla.add_row()
            _set_celda_by_cell(row.cells[0], "Nivel del Riesgo", bold=True, bg="D9E2F3")
            riesgo_valor = hallazgo.probabilidad * hallazgo.impacto
            texto_riesgo = (
                f"Probabilidad: {hallazgo.probabilidad}/5, "
                f"Impacto: {hallazgo.impacto}/5\n"
                f"(P={hallazgo.probabilidad}, I={hallazgo.impacto}, "
                f"R={riesgo_valor}) "
                f"Riesgo {hallazgo.nivel_riesgo}."
            )
            _set_celda_by_cell(row.cells[1], texto_riesgo)

        # ── Guardar ──────────────────────────────────────────────
        nombre_archivo = f"Fichas_Hallazgos_{contexto.nombre_entidad.replace(' ', '_')}.docx"
        ruta_salida = SALIDAS_DIR / nombre_archivo
        doc.save(str(ruta_salida))
        print(f"  ✅ Fichas de hallazgos generadas: {ruta_salida.name}")
        return ruta_salida

    except Exception as e:
        print(f"  ❌ Error generando fichas: {e}")
        return None


def generar_guia_pruebas(
    hallazgos: list[Hallazgo],
    contexto: ContextoEntidad,
) -> Path | None:
    """
    Genera el documento Word de Guía de Pruebas (Identificación de Evidencias).

    Args:
        hallazgos: Lista de hallazgos con evidencias.
        contexto: Contexto de la entidad.

    Returns:
        Ruta al archivo generado.
    """
    try:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

        # ── Encabezado ────────────────────────────────────────────
        p = doc.add_paragraph()
        run = p.add_run("AUDITORES: ")
        run.bold = True
        grupo = contexto.grupo_auditor or "Grupo Auditor"
        p.add_run(grupo).bold = True

        if contexto.integrantes:
            p = doc.add_paragraph()
            run = p.add_run("Integrantes: ")
            run.bold = True
            run.add_break()
            for nombre in contexto.integrantes:
                p.add_run(f"{nombre}\n")

        p = doc.add_paragraph()
        run = p.add_run("BANCO AUDITADO: ")
        run.bold = True
        p.add_run(contexto.nombre_entidad).bold = True

        # Título
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("IDENTIFICACIÓN DE EVIDENCIAS EN BASE A LOS HALLAZGOS")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 51, 102)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("GUÍA DE PRUEBAS")
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph()

        # ── Una sección por hallazgo ─────────────────────────────
        prefijo = _generar_prefijo_entidad(contexto.nombre_entidad)

        for i, hallazgo in enumerate(hallazgos):
            if i > 0:
                doc.add_page_break()

            codigo_guia = f"{prefijo}-{hallazgo.codigo}"

            # Subtítulo de la guía
            p = doc.add_paragraph()
            run = p.add_run(f"Guía de Pruebas")
            run.bold = True
            run.font.size = Pt(11)

            # Tabla de información
            tabla_info = doc.add_table(rows=5, cols=2)
            tabla_info.style = "Table Grid"

            _set_celda(tabla_info, 0, 0, codigo_guia, bold=True, bg="003366", fg="FFFFFF")
            _set_celda(tabla_info, 0, 1, "", bg="003366")

            _set_celda(tabla_info, 1, 0, "Dominio", bold=True, bg="D9E2F3")
            _set_celda(tabla_info, 1, 1, hallazgo.dominio)

            _set_celda(tabla_info, 2, 0, "Proceso", bold=True, bg="D9E2F3")
            _set_celda(tabla_info, 2, 1, "\n".join(hallazgo.procesos))

            _set_celda(tabla_info, 3, 0, "Objetivo de Control", bold=True, bg="D9E2F3")
            _set_celda(tabla_info, 3, 1, hallazgo.objetivo_control)

            _set_celda(tabla_info, 4, 0, "Riesgos Asociados", bold=True, bg="D9E2F3")
            _set_celda(tabla_info, 4, 1, ", ".join(hallazgo.riesgos_asociados))

            doc.add_paragraph()

            # Tabla de evidencias
            if hallazgo.evidencias:
                num_evidencias = len(hallazgo.evidencias)
                tabla_ev = doc.add_table(rows=num_evidencias + 1, cols=3)
                tabla_ev.style = "Table Grid"

                # Header
                _set_celda(tabla_ev, 0, 0, "No", bold=True, bg="003366", fg="FFFFFF")
                _set_celda(tabla_ev, 0, 1, "Evidencia", bold=True, bg="003366", fg="FFFFFF")
                _set_celda(tabla_ev, 0, 2, "Descripción", bold=True, bg="003366", fg="FFFFFF")

                for j, ev in enumerate(hallazgo.evidencias):
                    _set_celda(tabla_ev, j + 1, 0, str(ev.numero))
                    _set_celda(tabla_ev, j + 1, 1, ev.referencia)
                    _set_celda(tabla_ev, j + 1, 2, ev.descripcion)
            else:
                # Si no hay evidencias, crear tabla con citas
                citas = hallazgo.citas_bibliograficas
                tabla_ev = doc.add_table(rows=len(citas) + 1, cols=3)
                tabla_ev.style = "Table Grid"

                _set_celda(tabla_ev, 0, 0, "No", bold=True, bg="003366", fg="FFFFFF")
                _set_celda(tabla_ev, 0, 1, "Evidencia", bold=True, bg="003366", fg="FFFFFF")
                _set_celda(tabla_ev, 0, 2, "Descripción", bold=True, bg="003366", fg="FFFFFF")

                for j, cita in enumerate(citas):
                    ref = f"{cita.documento}, {cita.seccion}"
                    if cita.pagina:
                        ref += f", {cita.pagina}"
                    _set_celda(tabla_ev, j + 1, 0, str(j + 1))
                    _set_celda(tabla_ev, j + 1, 1, ref)
                    _set_celda(tabla_ev, j + 1, 2, cita.descripcion)

        # ── Guardar ──────────────────────────────────────────────
        nombre_archivo = f"Guia_Pruebas_{contexto.nombre_entidad.replace(' ', '_')}.docx"
        ruta_salida = SALIDAS_DIR / nombre_archivo
        doc.save(str(ruta_salida))
        print(f"  ✅ Guía de pruebas generada: {ruta_salida.name}")
        return ruta_salida

    except Exception as e:
        print(f"  ❌ Error generando guía de pruebas: {e}")
        return None


# ============================================================
# FUNCIONES AUXILIARES
# ============================================================

def _generar_prefijo_entidad(nombre: str) -> str:
    """Genera un prefijo corto a partir del nombre de la entidad (ej: 'FB' para 'Fondo de Bikini')."""
    palabras = nombre.split()
    if len(palabras) >= 2:
        # Tomar las iniciales de las primeras 2 palabras significativas
        significativas = [p for p in palabras if len(p) > 2]
        if len(significativas) >= 2:
            return (significativas[0][0] + significativas[1][0]).upper()
    return nombre[:2].upper()


def _set_celda(tabla, fila: int, col: int, texto: str, bold: bool = False,
               bg: str | None = None, fg: str | None = None):
    """Establece texto y formato de una celda de tabla."""
    cell = tabla.cell(fila, col)
    _set_celda_by_cell(cell, texto, bold, bg, fg)


def _set_celda_by_cell(cell, texto: str, bold: bool = False,
                       bg: str | None = None, fg: str | None = None):
    """Establece texto y formato de una celda por referencia directa."""
    cell.text = texto
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            if bold:
                run.bold = True
            if fg:
                run.font.color.rgb = RGBColor(
                    int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16)
                )

    if bg:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), bg)
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)
